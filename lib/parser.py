import io
import re
from typing import Optional

import pdfplumber
from pdfplumber.utils.exceptions import PdfminerException
from docx import Document


WORD_GAP_THRESHOLD = 1.5
LINE_GAP_THRESHOLD = 3.0


def _group_chars_into_lines(chars: list[dict]) -> str:
    """Group characters into lines and words, inserting spaces at word boundaries."""
    if not chars:
        return ""
    chars_sorted = sorted(chars, key=lambda c: (round(c["top"], 0), c["x0"]))
    lines: list[list[str]] = []
    current_line: list[str] = []
    current_word: list[str] = []
    prev_char: dict | None = None
    for c in chars_sorted:
        if prev_char and abs(c["top"] - prev_char["top"]) > LINE_GAP_THRESHOLD:
            if current_word:
                current_line.append("".join(current_word))
                current_word = []
            if current_line:
                lines.append(current_line)
                current_line = []
        if prev_char and abs(c["top"] - prev_char["top"]) <= LINE_GAP_THRESHOLD:
            gap = c["x0"] - prev_char["x1"]
            if gap > WORD_GAP_THRESHOLD:
                if current_word:
                    current_line.append("".join(current_word))
                    current_word = []
        current_word.append(c["text"])
        prev_char = c
    if current_word:
        current_line.append("".join(current_word))
    if current_line:
        lines.append(current_line)
    return "\n".join(" ".join(line) for line in lines)


def parse_pdf(file_bytes: bytes) -> str:
    try:
        page_texts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                line_text = _group_chars_into_lines(page.chars)
                if line_text:
                    page_texts.append(line_text)
        text = "\n".join(page_texts)
        text = re.sub(r"[^\S\n]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip() if text.strip() else ""
    except PdfminerException:
        return ""


def parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def parse_resume(file_bytes: bytes, filename: str) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}. Expected PDF or DOCX.")


def extract_email(text: str) -> Optional[str]:
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    patterns = [
        r"\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}",
        r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None


def has_images(file_bytes: bytes, filename: str) -> bool:
    if not filename.lower().endswith(".pdf"):
        return False
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                if page.images:
                    return True
    except PdfminerException:
        return False
    return False


def count_pages(file_bytes: bytes, filename: str) -> int:
    if filename.lower().endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return len(pdf.pages)
        except PdfminerException:
            return 1
    return 1


def has_tables(file_bytes: bytes, filename: str) -> bool:
    if not filename.lower().endswith(".pdf"):
        return False
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                if page.find_tables():
                    return True
    except PdfminerException:
        return False
    return False


def has_columns(file_bytes: bytes, filename: str) -> bool:
    if not filename.lower().endswith(".pdf"):
        return False
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                chars = page.chars
                if not chars:
                    continue
                x_positions = sorted(set(round(c["x0"]) for c in chars))
                distinct_columns = 1
                for i in range(1, len(x_positions)):
                    if x_positions[i] - x_positions[i - 1] > 50:
                        distinct_columns += 1
                if distinct_columns > 2:
                    return True
    except PdfminerException:
        return False
    return False


def count_words(text: str) -> int:
    return len(text.split())
