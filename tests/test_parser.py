import io
import pytest
from lib.parser import extract_email, extract_phone, count_words


SAMPLE_RESUME = """John Doe
john.doe@example.com
(555) 123-4567

SUMMARY
Experienced software engineer with 8 years of experience in full-stack development.

EXPERIENCE
Senior Software Engineer, Tech Corp
Jan 2020 - Present
- Led a team of 5 engineers to deliver a new payment platform
- Increased system throughput by 40% through performance optimization
- Reduced infrastructure costs by $120,000 annually

Software Engineer, Startup Inc
Mar 2017 - Dec 2019
- Developed REST APIs serving 100,000 daily active users
- Implemented CI/CD pipeline reducing deployment time by 70%

EDUCATION
BS Computer Science, University of Technology, 2015

SKILLS
Python, JavaScript, React, Node.js, PostgreSQL, AWS, Docker, Kubernetes
"""


def test_extract_email():
    assert extract_email(SAMPLE_RESUME) == "john.doe@example.com"
    assert extract_email("No email here") is None
    assert extract_email("contact@company.co.uk") == "contact@company.co.uk"


def test_extract_phone():
    assert extract_phone(SAMPLE_RESUME) == "(555) 123-4567"
    assert extract_phone("No phone here") is None
    assert extract_phone("Call +1-555-123-4567 anytime") == "+1-555-123-4567"
    assert extract_phone("555.123.4567") == "555.123.4567"


def test_count_words():
    assert count_words("") == 0
    assert count_words("Hello world") == 2
    assert count_words("One two three four five") == 5
    assert count_words(SAMPLE_RESUME) > 20


def test_parse_pdf_from_bytes():
    from lib.parser import parse_pdf
    result = parse_pdf(b"%PDF-1.4\n%invalid")
    assert isinstance(result, str)
    assert result == ""


def test_parse_docx_structure():
    from docx import Document
    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Test resume content")
    doc.add_paragraph("Skills: Python, React")
    doc.save(buffer)
    buffer.seek(0)

    from lib.parser import parse_docx
    result = parse_docx(buffer.read())
    assert "Test resume content" in result
    assert "Python, React" in result


def test_parse_resume_unsupported_format():
    from lib.parser import parse_resume
    with pytest.raises(ValueError, match="Unsupported file format"):
        parse_resume(b"data", "resume.txt")


def test_parse_resume_pdf_invalid():
    from lib.parser import parse_resume, parse_pdf
    result = parse_resume(b"%PDF-1.4\ninvalid", "resume.pdf")
    assert isinstance(result, str)
    assert result == ""
